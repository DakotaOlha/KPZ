from datetime import datetime
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
# Додайте цей імпорт, якщо його немає, для вирівнювання
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DataExporter:

    def __init__(self, db_manager):
        self.db = db_manager

    def export_to_word(self):
        try:
            words = self.db.get_all_words()

            if not words:
                messagebox.showwarning("Експорт", "Немає слів для експорту!")
                return

            doc = Document()

            title = doc.add_heading('Звіт - Статистика навчання', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Дата створення: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_run.italic = True
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            cat_stats = self.db.get_category_statistics()

            if cat_stats:
                doc.add_heading('Статистика по категоріям', level=1)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'

                hdr_cells = table.rows[0].cells
                headers = ['Категорія', 'Всього', 'Вивчено', 'Вивчається', 'Нові']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)

                for cat in cat_stats:
                    row_cells = table.add_row().cells
                    row_cells[0].text = cat[0] or "Без категорії"
                    row_cells[1].text = str(cat[1])
                    row_cells[2].text = str(cat[2] or 0)
                    row_cells[3].text = str(cat[3] or 0)
                    row_cells[4].text = str(cat[4] or 0)
            else:
                doc.add_paragraph("Немає даних по категоріях для відображення.")

            section = doc.sections[0]
            footer = section.footer

            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            LOGO_PATH = 'logo.png'

            try:
                run_logo = footer_para.add_run()
                run_logo.add_picture(LOGO_PATH, height=Inches(0.3))
                footer_para.add_run("   ")
            except FileNotFoundError:
                print(f"Warning: Logo file not found at '{LOGO_PATH}'. Skipping logo.")
            except Exception as e:
                print(f"Warning: Could not add logo. Error: {e}")

            stats = self.db.get_statistics()
            stats_text = (
                f"Загальна статистика: "
                f"Всього слів: {stats['total_words']} | "
                f"Вивчено: {stats['learned_words']} | "
                f"Вивчається: {stats['learning_words']} | "
                f"Нові: {stats['new_words']} | "
                f"Прогрес: {stats['progress_percentage']:.1f}%"
            )

            run_stats = footer_para.add_run(stats_text)
            run_stats.font.size = Pt(9)
            run_stats.font.italic = True

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=f"LearnEasy_Statistics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Успіх", f"Файл експортовано:\n{file_path}")
                return file_path

        except Exception as e:
            messagebox.showerror("Помилка експорту", f"Не вдалося експортувати:\n{str(e)}")

    def export_to_excel(self):
        try:
            words = self.db.get_all_words()

            if not words:
                messagebox.showwarning("Експорт", "Немає слів для експорту!")
                return

            wb = Workbook()
            ws = wb.active
            ws.title = "Слова"

            headers = ['№', 'Слово', 'Переклад', 'Транскрипція', 'Рівень', 'Категорія',
                       'Показів', 'Правильно', 'Неправильно', 'Улюблене', 'Складність']
            ws.append(headers)

            header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)

            for col in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=col)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")

            for idx, word in enumerate(words, 1):
                ws.append([
                    idx,
                    word[1],
                    word[2],
                    word[3] or "",
                    word[4],
                    word[5] or "-",
                    word[6],
                    word[7],
                    word[8],
                    "✓" if word[9] else "",
                    word[10]
                ])

            column_widths = [5, 20, 20, 20, 10, 15, 10, 12, 12, 10, 12]
            for i, width in enumerate(column_widths, 1):
                ws.column_dimensions[chr(64 + i)].width = width

            stats_ws = wb.create_sheet("Статистика")
            stats = self.db.get_statistics()

            stats_ws['A1'] = "Загальна статистика"
            stats_ws['A1'].font = Font(bold=True, size=14)

            stats_data = [
                ["Показник", "Значення"],
                ["Всього слів", stats['total_words']],
                ["Вивчено", stats['learned_words']],
                ["Вивчається", stats['learning_words']],
                ["Нові", stats['new_words']],
                ["Прогрес, %", f"{stats['progress_percentage']:.1f}"],
                ["Дата звіту", datetime.now().strftime('%d.%m.%Y %H:%M')]
            ]

            for row in stats_data:
                stats_ws.append(row)

            stats_ws.column_dimensions['A'].width = 25
            stats_ws.column_dimensions['B'].width = 20

            stat_header_fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
            stat_header_font = Font(bold=True, color="FFFFFF")

            for col in range(1, 3):
                cell = stats_ws.cell(row=1, column=col)
                cell.fill = stat_header_fill
                cell.font = stat_header_font

            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")],
                initialfile=f"LearnEasy_Words_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )

            if file_path:
                wb.save(file_path)
                messagebox.showinfo("Успіх", f"Файл експортовано:\n{file_path}")
                return file_path

        except Exception as e:
            messagebox.showerror("Помилка експорту", f"Не вдалося експортувати:\n{str(e)}")

    def export_statistics_to_word(self):
        try:
            doc = Document()

            title = doc.add_heading('Звіт - Статистика навчання', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Дата створення: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_run.italic = True
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            cat_stats = self.db.get_category_statistics()

            if cat_stats:
                doc.add_heading('Статистика по категоріям', level=1)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'

                hdr_cells = table.rows[0].cells
                headers = ['Категорія', 'Всього', 'Вивчено', 'Вивчається', 'Нові']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)

                for cat in cat_stats:
                    row_cells = table.add_row().cells
                    row_cells[0].text = cat[0] or "Без категорії"
                    row_cells[1].text = str(cat[1])
                    row_cells[2].text = str(cat[2] or 0)
                    row_cells[3].text = str(cat[3] or 0)
                    row_cells[4].text = str(cat[4] or 0)
            else:
                doc.add_paragraph("Немає даних по категоріях для відображення.")

            section = doc.sections[0]
            footer = section.footer

            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            LOGO_PATH = 'logo.png'

            try:
                run_logo = footer_para.add_run()
                run_logo.add_picture(LOGO_PATH, height=Inches(0.3))
                footer_para.add_run("   ")
            except FileNotFoundError:
                print(f"Warning: Logo file not found at '{LOGO_PATH}'. Skipping logo.")
            except Exception as e:
                print(f"Warning: Could not add logo. Error: {e}")

            stats = self.db.get_statistics()
            stats_text = (
                f"Загальна статистика: "
                f"Всього слів: {stats['total_words']} | "
                f"Вивчено: {stats['learned_words']} | "
                f"Вивчається: {stats['learning_words']} | "
                f"Нові: {stats['new_words']} | "
                f"Прогрес: {stats['progress_percentage']:.1f}%"
            )

            run_stats = footer_para.add_run(stats_text)
            run_stats.font.size = Pt(9)
            run_stats.font.italic = True

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=f"LearnEasy_Statistics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Успіх", f"Файл експортовано:\n{file_path}")
                return file_path

        except Exception as e:
            messagebox.showerror("Помилка експорту", f"Не вдалося експортувати:\n{str(e)}")